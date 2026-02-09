"""
Document converter service for PDF and Word documents.
Converts PDF/Word to images for OCR processing.
"""
import os
import tempfile
from pathlib import Path
from typing import Optional


def is_pdf(filename: str) -> bool:
    """Check if file is a PDF"""
    return filename.lower().endswith('.pdf')


def is_word(filename: str) -> bool:
    """Check if file is a Word document"""
    ext = filename.lower().split('.')[-1] if '.' in filename else ''
    return ext in ('doc', 'docx')


def convert_pdf_to_images(pdf_path: str, output_dir: Optional[str] = None) -> list[str]:
    """
    Convert PDF pages to images.
    Returns list of image file paths.
    """
    try:
        from pdf2image import convert_from_path
    except ImportError:
        raise ImportError("pdf2image не установлен. Установите: pip install pdf2image")

    if output_dir is None:
        output_dir = tempfile.mkdtemp()

    # Convert PDF to images
    images = convert_from_path(pdf_path, dpi=200)

    image_paths = []
    for i, image in enumerate(images):
        image_path = os.path.join(output_dir, f"page_{i + 1}.png")
        image.save(image_path, "PNG")
        image_paths.append(image_path)

    return image_paths


def convert_word_to_images(word_path: str, output_dir: Optional[str] = None) -> list[str]:
    """
    Convert Word document to images.
    First converts to PDF, then to images.
    For simplicity, extracts text and creates a text file for OCR.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx не установлен. Установите: pip install python-docx")

    if output_dir is None:
        output_dir = tempfile.mkdtemp()

    # Extract text from Word document
    doc = Document(word_path)

    # Collect all text
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text)
            full_text.append(" | ".join(row_text))

    text_content = "\n".join(full_text)

    # Create a simple image with text for OCR
    # Or return a text file that can be processed
    try:
        from PIL import Image, ImageDraw, ImageFont

        # Create image with text
        lines = text_content.split('\n')

        # Calculate image size
        line_height = 20
        margin = 50
        width = 1200
        height = max(800, len(lines) * line_height + margin * 2)

        # Create white image
        img = Image.new('RGB', (width, height), color='white')
        draw = ImageDraw.Draw(img)

        # Try to use a system font
        try:
            font = ImageFont.truetype("arial.ttf", 14)
        except:
            font = ImageFont.load_default()

        # Draw text
        y = margin
        for line in lines:
            if line.strip():
                draw.text((margin, y), line, fill='black', font=font)
            y += line_height

        image_path = os.path.join(output_dir, "word_page_1.png")
        img.save(image_path, "PNG")
        return [image_path]

    except Exception as e:
        # Fallback: save as text file
        text_path = os.path.join(output_dir, "word_content.txt")
        with open(text_path, 'w', encoding='utf-8') as f:
            f.write(text_content)
        return [text_path]


def extract_text_from_word(word_path: str) -> str:
    """
    Extract plain text from Word document.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx не установлен. Установите: pip install python-docx")

    doc = Document(word_path)

    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text)
            full_text.append(" | ".join(row_text))

    return "\n".join(full_text)


def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extract plain text from PDF using PyPDF2.
    """
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise ImportError("PyPDF2 не установлен. Установите: pip install pypdf2")

    reader = PdfReader(pdf_path)

    text_parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)

    return "\n".join(text_parts)


def convert_document_for_ocr(file_path: str, output_dir: Optional[str] = None) -> dict:
    """
    Convert document to format suitable for OCR.
    Returns dict with:
    - images: list of image paths (for OCR)
    - extracted_text: text extracted directly (for Word/PDF)
    - original_type: 'pdf', 'word', or 'image'
    """
    filename = os.path.basename(file_path)

    if is_pdf(filename):
        return {
            "images": convert_pdf_to_images(file_path, output_dir),
            "extracted_text": extract_text_from_pdf(file_path),
            "original_type": "pdf"
        }
    elif is_word(filename):
        return {
            "images": convert_word_to_images(file_path, output_dir),
            "extracted_text": extract_text_from_word(file_path),
            "original_type": "word"
        }
    else:
        # Regular image - no conversion needed
        return {
            "images": [file_path],
            "extracted_text": None,
            "original_type": "image"
        }
